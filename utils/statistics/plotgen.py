from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any, Iterable

import matplotlib.pyplot as plt
import numpy as np
from scipy.interpolate import CubicSpline

from utils.io.xyz_io import import_multiframe_xyz


HARTREE_TO_KCALMOL = 627.5094740631


try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import (
        WD_CELL_VERTICAL_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_TABLE_ALIGNMENT,
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ======================================================================================
# STATE / NAMING HELPERS
# ======================================================================================

KNOWN_GUESS_LABELS = {
    "wave1basic",
    "wave1concoursed",
    "wave2basic",
    "wave2concoursed",
}


def _dedupe(items: Iterable[Any]) -> list[str]:
    out = []
    seen = set()

    for item in items:
        if item is None:
            continue

        item = str(item)

        if item in seen:
            continue

        seen.add(item)
        out.append(item)

    return out


def _bool_state(state: dict, *keys: str, default: bool = False) -> bool:
    for key in keys:
        if key in state:
            return bool(state[key])

    return default


def _all_neb_muted(state: dict) -> bool:
    return _bool_state(
        state,
        "mute_all_neb",
        "prepare_guesses_only",
        default=False,
    )


def _standard_neb_muted(state: dict) -> bool:
    return bool(
        _bool_state(
            state,
            "mute_standard_neb",
            "mute_standart_neb",
            default=False,
        )
        or _all_neb_muted(state)
    )


def _concours_enabled(state: dict) -> bool:
    return bool(
        state.get("reparam_scf_check", True)
        and not state.get("mute_concours", False)
    )


def _guess_suffix(state: dict) -> str:
    return "concoursed" if _concours_enabled(state) else "basic"


def _normalize_guess_method_name(method_name: str, state: dict) -> str:
    method_name = str(method_name)

    if method_name in {"wave1", "wave2"}:
        return f"{method_name}{_guess_suffix(state)}"

    if method_name in KNOWN_GUESS_LABELS:
        return method_name

    return method_name


def _standard_run_name_candidates_from_state(state: dict) -> list[str]:
    out = []

    if state.get("standard_neb_run_name"):
        out.append(str(state["standard_neb_run_name"]))

    # New compact naming.
    out.append("standard")

    if state.get("config_hash"):
        out.append(f"standard_{state['config_hash']}")

    # Old/native naming.
    if "base_name" in state and "pair_tag" in state:
        out.append(f"{state['base_name']}_standard_{state['pair_tag']}")

    if "base_name" in state and "pair_tag" in state:
        h = state.get("config_hash", "nohash")
        out.append(f"{state['base_name']}_standard_{state['pair_tag']}_{h}")

    return _dedupe(out)


def _custom_run_name_candidates_from_state(state: dict, method_name: str) -> list[str]:
    method_name = _normalize_guess_method_name(method_name, state)

    out = [
        method_name,
    ]

    if "base_name" in state and "pair_tag" in state:
        out.append(f"{state['base_name']}_myguess_{method_name}_{state['pair_tag']}")

    return _dedupe(out)


def _run_paths(neb_folder: Path, run_name: str) -> dict[str, Path]:
    run_dir = Path(neb_folder) / f"neb_{run_name}"

    return {
        "run_dir": run_dir,
        "npz": run_dir / f"{run_name}_summary.npz",
        "json": run_dir / f"{run_name}_results.json",
        "xyz": run_dir / f"{run_name}_final_chain.xyz",
    }


def _run_dir_exists(neb_folder: Path, run_name: str) -> bool:
    return (Path(neb_folder) / f"neb_{run_name}").exists()


def _run_has_outputs(neb_folder: Path, run_name: str) -> bool:
    paths = _run_paths(Path(neb_folder), run_name)

    return bool(
        paths["run_dir"].exists()
        and (
            paths["npz"].exists()
            or paths["json"].exists()
        )
    )


def _pick_existing_run_name(neb_folder: Path, candidates: list[str]) -> str | None:
    neb_folder = Path(neb_folder)

    # Prefer finished runs.
    for candidate in candidates:
        if _run_has_outputs(neb_folder, candidate):
            return candidate

    # Then accept existing run folders.
    for candidate in candidates:
        if _run_dir_exists(neb_folder, candidate):
            return candidate

    return candidates[0] if candidates else None


def _custom_run_names_from_state(state: dict) -> list[str]:
    neb_folder = Path(state["base_cfg"].neb_folder)

    if state.get("custom_neb_run_names"):
        return _dedupe(str(x) for x in state["custom_neb_run_names"])

    if state.get("final_guess_npzs"):
        methods = list(state["final_guess_npzs"].keys())
    else:
        methods = list(state.get("custom_guess_methods", []))

    out = []

    for method in methods:
        candidates = _custom_run_name_candidates_from_state(state, method)
        picked = _pick_existing_run_name(neb_folder, candidates)

        if picked is not None:
            out.append(picked)

    return _dedupe(out)


def _standard_run_name_from_state(state: dict) -> str | None:
    if _standard_neb_muted(state):
        return None

    neb_folder = Path(state["base_cfg"].neb_folder)

    return _pick_existing_run_name(
        neb_folder,
        _standard_run_name_candidates_from_state(state),
    )


def _out_stem_from_state(state: dict) -> str:
    return str(state.get("name", state.get("pair_tag", "neb")))


# ======================================================================================
# PUBLIC API
# ======================================================================================

def compare_all_neb(*args, force: bool = True, **kwargs):
    """
    Supports two call styles:

        compare_all_neb(state)

    or:

        compare_all_neb(
            neb_folder,
            run_standard,
            custom_runs,
            out_stem,
            force=True,
        )

    Creates in neb_folder:
      - *_all_neb_profiles.png
      - *_neb_report_plot.png
      - *_all_neb_table.csv
      - *_all_neb_table.txt
      - *_neb_report.docx
    """
    if len(args) == 1 and isinstance(args[0], dict):
        return _compare_all_neb_from_state(args[0], force=force)

    return _compare_all_neb_explicit(*args, force=force, **kwargs)


def generate_word_report(final_states, out_docx: str | Path = "Full_Project_Summary.docx", force: bool = True):
    """
    Generates a master DOCX report for:
      - one normal state dict;
      - list of states;
      - FINAL_STATE = {"multi": True, "states": {...}}.

    Also ensures each segment/job gets its own one-page DOCX in its neb folder.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Install it with:\n"
            "    pip install python-docx"
        )

    states = list(_iter_report_states(final_states))

    if not states:
        print("No states found for master DOCX report.")
        return None

    doc = Document()
    _setup_landscape_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Full Project NEB Summary")
    title_run.bold = True
    title_run.font.size = Pt(15)

    added_any = False
    generated_local_docs = []

    for state in states:
        if "base_cfg" not in state:
            continue

        if _all_neb_muted(state):
            print(f"Skipping DOCX for {state.get('name', state.get('pair_tag', 'unknown'))}: all NEB muted.")
            continue

        result = compare_all_neb(state, force=force)

        if not result:
            continue

        csv_path = result.get("csv")
        plot_path = result.get("docx_plot") or result.get("plot")
        local_docx = result.get("docx")

        if not csv_path or not Path(csv_path).exists():
            print(f"Skipping master DOCX block: missing CSV table for {state.get('name')}")
            continue

        if added_any:
            doc.add_page_break()

        _add_state_to_master_doc(
            doc=doc,
            state=state,
            csv_path=Path(csv_path),
            plot_path=Path(plot_path) if plot_path else None,
        )

        if local_docx:
            generated_local_docs.append(Path(local_docx))

        added_any = True

    if not added_any:
        print("No finished NEB reports were added to master DOCX.")
        return None

    out_docx = Path(out_docx)
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)

    print(f"\nSaved master DOCX report: {out_docx}")

    if generated_local_docs:
        print("\nSaved local one-page DOCX reports:")
        for path in generated_local_docs:
            print(f"  {path}")

    return out_docx


# ======================================================================================
# COMPARE ALL NEB
# ======================================================================================

def _compare_all_neb_from_state(state: dict, force: bool = True):
    if _all_neb_muted(state):
        print("All NEB runs are muted. Skipping NEB comparison report.")
        return None

    neb_folder = Path(state["base_cfg"].neb_folder)
    out_stem = _out_stem_from_state(state)

    run_standard = _standard_run_name_from_state(state)
    custom_runs = _custom_run_names_from_state(state)

    return _compare_all_neb_explicit(
        neb_folder=neb_folder,
        run_standard=run_standard,
        custom_runs=custom_runs,
        out_stem=out_stem,
        force=force,
    )


def _compare_all_neb_explicit(
    neb_folder,
    run_standard=None,
    custom_runs=None,
    out_stem: str = "neb_comparison",
    force: bool = True,
):
    neb_folder = Path(neb_folder)
    custom_runs = list(custom_runs or [])

    runs = []

    if run_standard is not None:
        runs.append(str(run_standard))

    runs.extend(str(x) for x in custom_runs if x is not None)
    runs = _dedupe(runs)

    if not runs:
        print("No NEB runs requested for comparison.")
        return None

    neb_folder.mkdir(parents=True, exist_ok=True)

    plot_path = neb_folder / f"{out_stem}_all_neb_profiles.png"
    docx_plot_path = neb_folder / f"{out_stem}_neb_report_plot.png"
    csv_path = neb_folder / f"{out_stem}_all_neb_table.csv"
    txt_path = neb_folder / f"{out_stem}_all_neb_table.txt"
    docx_path = neb_folder / f"{out_stem}_neb_report.docx"

    if (
        not force
        and plot_path.exists()
        and docx_plot_path.exists()
        and csv_path.exists()
        and txt_path.exists()
        and (docx_path.exists() or not DOCX_AVAILABLE)
    ):
        print("All-NEB comparison already exists. Skipping generation.")
        print(txt_path.read_text(encoding="utf-8"))

        return {
            "plot": plot_path,
            "docx_plot": docx_plot_path,
            "csv": csv_path,
            "txt": txt_path,
            "docx": docx_path if docx_path.exists() else None,
            "rows": _read_csv_rows(csv_path),
        }

    rows = []
    plot_items = []

    for run_name in runs:
        paths = _run_paths(neb_folder, run_name)
        label = _run_label(run_name, run_standard)

        if not paths["run_dir"].exists():
            print(f"Warning: run folder not found: {paths['run_dir']}")
            continue

        if not paths["npz"].exists() and not paths["json"].exists():
            print(f"Warning: run has no summary/results yet: {paths['run_dir']}")
            continue

        json_data = _load_json(paths["json"])
        plot_data = _extract_plot_data(paths["npz"], paths["xyz"])
        interp_x, interp_e = _interpolated_peak(plot_data)

        final_activation = _get_from_json(
            json_data,
            "final",
            "activation",
            HARTREE_TO_KCALMOL,
        )

        initial_activation = _get_from_json(
            json_data,
            "initial",
            "activation",
            HARTREE_TO_KCALMOL,
        )

        try:
            rxn_energy = (
                json_data["final"]["end"] - json_data["final"]["start"]
            ) * HARTREE_TO_KCALMOL
        except Exception:
            rxn_energy = float(plot_data["energies"][-1]) if plot_data is not None else None

        row = {
            "run": label,
            "run_name": run_name,
            "converged": _get_from_json(json_data, "converged"),
            "optimizer_steps": _get_from_json(json_data, "nsteps_taken"),
            "neb_cycles": _get_from_json(json_data, "n_neb_cycles"),
            "initial_activation_kcalmol": initial_activation,
            "initial_peak_index": _get_from_json(json_data, "initial", "peak_index"),
            "initial_fmax_ev_ang": _get_from_json(json_data, "initial", "neb_fmax"),
            "final_activation_kcalmol": final_activation,
            "final_peak_index": _get_from_json(json_data, "final", "peak_index"),
            "final_fmax_ev_ang": _get_from_json(json_data, "final", "neb_fmax"),
            "products_minus_reactants_kcalmol": rxn_energy,
            "interp_peak_x_0_1": interp_x,
            "interp_activation_kcalmol": interp_e,
        }

        rows.append(row)

        if plot_data is not None:
            plot_items.append((label, plot_data))

    if not rows:
        print("No finished NEB runs found for comparison.")
        return None

    if plot_items:
        _save_full_png_report(
            plot_items=plot_items,
            out_stem=out_stem,
            plot_path=plot_path,
        )

        _save_docx_png_plot(
            plot_items=plot_items,
            out_stem=out_stem,
            plot_path=docx_plot_path,
        )

    _save_csv_table(csv_path, rows)
    txt = _save_txt_table(txt_path, rows, out_stem)

    if DOCX_AVAILABLE:
        _save_docx_report(
            docx_path=docx_path,
            plot_path=docx_plot_path if plot_items else None,
            rows=rows,
            out_stem=out_stem,
        )
    else:
        print("Warning: python-docx is not installed. DOCX report was not created.")

    print(txt)

    if plot_items:
        print(f"\nSaved plot: {plot_path}")
        print(f"Saved DOCX plot: {docx_plot_path}")

    print(f"Saved CSV table: {csv_path}")
    print(f"Saved TXT table: {txt_path}")

    if DOCX_AVAILABLE:
        print(f"Saved DOCX report: {docx_path}")

    return {
        "plot": plot_path if plot_items else None,
        "docx_plot": docx_plot_path if plot_items else None,
        "csv": csv_path,
        "txt": txt_path,
        "docx": docx_path if DOCX_AVAILABLE else None,
        "rows": rows,
    }


# ======================================================================================
# DATA EXTRACTION
# ======================================================================================

def _run_label(run_name: str, run_standard: str | None = None) -> str:
    if run_standard is not None and run_name == run_standard:
        return "standard"

    if run_name in KNOWN_GUESS_LABELS:
        return run_name

    if "_myguess_" in run_name:
        label = run_name.split("_myguess_", 1)[1]

        for known in [
            "wave1basic",
            "wave1concoursed",
            "wave2basic",
            "wave2concoursed",
        ]:
            if label.startswith(known):
                return known

        if "_" in label:
            label = label.rsplit("_", 1)[0]

        return label

    return run_name


def _load_json(path: Path) -> dict:
    path = Path(path)

    if not path.exists():
        return {}

    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def _get_from_json(data: dict, key1: str, key2: str | None = None, mult: float = 1.0):
    try:
        value = data.get(key1, {})

        if key2 is not None:
            value = value.get(key2)

        if isinstance(value, (int, float)):
            return value * mult

        return value
    except Exception:
        return None


def _extract_plot_data(npz_path: Path, xyz_path: Path):
    npz_path = Path(npz_path)
    xyz_path = Path(xyz_path)

    if not npz_path.exists():
        return None

    with np.load(npz_path, allow_pickle=True) as data:
        if "final_rel_start" in data:
            energies = np.asarray(data["final_rel_start"], dtype=float) * HARTREE_TO_KCALMOL
        elif "final_rel" in data:
            energies = np.asarray(data["final_rel"], dtype=float) * HARTREE_TO_KCALMOL
            energies = energies - energies[0]
        else:
            print(f"Warning: no final_rel/final_rel_start in {npz_path}")
            return None

    bead_idx = np.arange(len(energies))

    if xyz_path.exists():
        _, geoms = import_multiframe_xyz(xyz_path)

        distances = [0.0]
        rmsds = [0.0]
        curr_dist = 0.0

        for i in range(1, len(geoms)):
            step_dist = float(np.sqrt(np.sum((geoms[i] - geoms[i - 1]) ** 2)))
            curr_dist += step_dist
            distances.append(curr_dist)

            rmsd = float(np.sqrt(np.mean(np.sum((geoms[i] - geoms[0]) ** 2, axis=1))))
            rmsds.append(rmsd)

        distances = np.asarray(distances[:len(energies)], dtype=float)
        rmsds = np.asarray(rmsds[:len(energies)], dtype=float)

        total_dist = distances[-1] if len(distances) and distances[-1] > 1e-12 else 1.0
        norm_distances = distances / total_dist
    else:
        distances = bead_idx.astype(float)
        norm_distances = bead_idx / max(len(bead_idx) - 1, 1)
        rmsds = np.full_like(bead_idx, np.nan, dtype=float)

    return {
        "energies": energies,
        "bead_idx": bead_idx,
        "distances": distances,
        "norm_distances": norm_distances,
        "rmsds": rmsds,
    }


def _interpolated_peak(plot_data):
    if plot_data is None:
        return None, None

    x = np.asarray(plot_data["norm_distances"], dtype=float)
    y = np.asarray(plot_data["energies"], dtype=float)

    try:
        x_unique, unique_idx = np.unique(x, return_index=True)
        y_unique = y[unique_idx]

        if len(x_unique) < 4:
            max_idx = int(np.argmax(y))
            return float(x[max_idx]), float(y[max_idx])

        cs = CubicSpline(x_unique, y_unique)
        x_dense = np.linspace(0.0, 1.0, 5000)
        y_dense = cs(x_dense)

        max_idx = int(np.argmax(y_dense))
        return float(x_dense[max_idx]), float(y_dense[max_idx])

    except Exception:
        max_idx = int(np.argmax(y))
        return float(x[max_idx]), float(y[max_idx])


# ======================================================================================
# PNG REPORTS
# ======================================================================================

def _save_full_png_report(plot_items, out_stem: str, plot_path: Path) -> None:
    plot_path = Path(plot_path)
    plot_path.parent.mkdir(parents=True, exist_ok=True)

    fig, axes = plt.subplots(1, 2, figsize=(14, 5))
    ax_idx, ax_norm = axes

    for label, data in plot_items:
        ax_idx.plot(
            data["bead_idx"],
            data["energies"],
            marker="o",
            linewidth=1.8,
            label=label,
        )

        ax_norm.plot(
            data["norm_distances"],
            data["energies"],
            marker="o",
            linewidth=1.8,
            label=label,
        )

    ax_idx.set_title("NEB profiles by bead index")
    ax_idx.set_xlabel("Bead index")
    ax_idx.set_ylabel("Relative energy, kcal/mol")
    ax_idx.grid(True, alpha=0.3)
    ax_idx.legend(fontsize=9)

    ax_norm.set_title("NEB profiles by normalized reaction coordinate")
    ax_norm.set_xlabel("Normalized reaction coordinate")
    ax_norm.set_ylabel("Relative energy, kcal/mol")
    ax_norm.grid(True, alpha=0.3)
    ax_norm.legend(fontsize=9)

    fig.suptitle(f"All NEB runs: {out_stem}", fontsize=14)
    fig.tight_layout()
    fig.savefig(plot_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


def _save_docx_png_plot(plot_items, out_stem: str, plot_path: Path) -> None:
    plot_path = Path(plot_path)
    plot_path.parent.mkdir(parents=True, exist_ok=True)

    fig, ax = plt.subplots(figsize=(10.8, 3.25))

    for label, data in plot_items:
        ax.plot(
            data["norm_distances"],
            data["energies"],
            marker="o",
            linewidth=1.7,
            markersize=4,
            label=label,
        )

    ax.set_title(f"NEB profiles: {out_stem}", fontsize=11)
    ax.set_xlabel("Normalized reaction coordinate", fontsize=9)
    ax.set_ylabel("Relative energy, kcal/mol", fontsize=9)
    ax.grid(True, alpha=0.3)
    ax.tick_params(axis="both", labelsize=8)

    ncol = min(max(len(plot_items), 1), 5)
    ax.legend(fontsize=7, ncol=ncol, loc="best")

    fig.tight_layout()
    fig.savefig(plot_path, dpi=220, bbox_inches="tight")
    plt.close(fig)


# ======================================================================================
# CSV / TXT TABLES
# ======================================================================================

def _save_csv_table(csv_path: Path, rows: list[dict]) -> None:
    csv_path = Path(csv_path)
    csv_path.parent.mkdir(parents=True, exist_ok=True)

    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()

        for row in rows:
            clean_row = {
                k: f"{v:.6f}" if isinstance(v, float) else v
                for k, v in row.items()
            }
            writer.writerow(clean_row)


def _read_csv_rows(csv_path: Path) -> list[dict]:
    csv_path = Path(csv_path)

    if not csv_path.exists():
        return []

    with csv_path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def _fmt(v) -> str:
    if v is None:
        return "N/A"

    if isinstance(v, str):
        if v.strip() == "":
            return "N/A"

        try:
            x = float(v)
            return f"{x:.3f}"
        except Exception:
            return v

    if isinstance(v, float):
        return f"{v:.3f}"

    return str(v)


def _report_metrics() -> list[tuple[str, str]]:
    return [
        ("Converged", "converged"),
        ("Optimizer steps", "optimizer_steps"),
        ("NEB cycles", "neb_cycles"),
        ("Initial Ea, kcal/mol", "initial_activation_kcalmol"),
        ("Initial peak index", "initial_peak_index"),
        ("Initial fmax, eV/A", "initial_fmax_ev_ang"),
        ("Final Ea, kcal/mol", "final_activation_kcalmol"),
        ("Final peak index", "final_peak_index"),
        ("Final fmax, eV/A", "final_fmax_ev_ang"),
        ("P - R, kcal/mol", "products_minus_reactants_kcalmol"),
        ("Spline peak x", "interp_peak_x_0_1"),
        ("Spline Ea, kcal/mol", "interp_activation_kcalmol"),
    ]


def _docx_metrics() -> list[tuple[str, str]]:
    return [
        ("Converged", "converged"),
        ("Steps", "optimizer_steps"),
        ("Cycles", "neb_cycles"),
        ("Final Ea", "final_activation_kcalmol"),
        ("Peak bead", "final_peak_index"),
        ("Final fmax", "final_fmax_ev_ang"),
        ("P - R", "products_minus_reactants_kcalmol"),
        ("Spline Ea", "interp_activation_kcalmol"),
    ]


def _save_txt_table(txt_path: Path, rows: list[dict], out_stem: str) -> str:
    txt_path = Path(txt_path)
    txt_path.parent.mkdir(parents=True, exist_ok=True)

    metrics = _report_metrics()

    col0 = 28
    col = 20
    width = col0 + (col + 3) * len(rows)

    lines = [
        "=" * width,
        f"ALL NEB COMPARISON TABLE: {out_stem}",
        "=" * width,
    ]

    header = f"{'Metric':<{col0}}" + "".join(
        f" | {row['run']:<{col}}"
        for row in rows
    )

    lines.append(header)
    lines.append("-" * len(header))

    for metric_name, key in metrics:
        line = f"{metric_name:<{col0}}" + "".join(
            f" | {_fmt(row.get(key)):<{col}}"
            for row in rows
        )
        lines.append(line)

    lines.append("=" * len(header))

    txt = "\n".join(lines)
    txt_path.write_text(txt, encoding="utf-8")

    return txt


# ======================================================================================
# DOCX LOCAL REPORT
# ======================================================================================

def _setup_landscape_doc(doc) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.25)
    section.right_margin = Inches(0.25)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(7)


def _save_docx_report(
    *,
    docx_path: Path,
    plot_path: Path | None,
    rows: list[dict],
    out_stem: str,
) -> None:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed. Install it with: pip install python-docx")

    docx_path = Path(docx_path)
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_landscape_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"NEB comparison report: {out_stem}")
    title_run.bold = True
    title_run.font.size = Pt(12)

    if plot_path is not None and Path(plot_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        run = p.add_run()
        run.add_picture(str(plot_path), width=Inches(10.35))

    _add_docx_table(
        doc=doc,
        rows=rows,
        compact=True,
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(2)
    footer_run = footer.add_run("Ea and P - R in kcal/mol; fmax in eV/A.")
    footer_run.font.size = Pt(6)

    doc.save(docx_path)


def _add_docx_table(doc, rows: list[dict], compact: bool = True) -> None:
    metrics = _docx_metrics() if compact else _report_metrics()

    table = doc.add_table(rows=len(metrics) + 1, cols=len(rows) + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    font_size = 6.5 if len(rows) <= 5 else 5.5

    header_cells = table.rows[0].cells
    _set_cell_text(header_cells[0], "Metric", bold=True, size=font_size)

    for j, row in enumerate(rows, start=1):
        _set_cell_text(header_cells[j], str(row["run"]), bold=True, size=font_size)

    for i, (metric_name, key) in enumerate(metrics, start=1):
        cells = table.rows[i].cells
        _set_cell_text(cells[0], metric_name, bold=True, size=font_size)

        for j, row in enumerate(rows, start=1):
            _set_cell_text(cells[j], _fmt(row.get(key)), bold=False, size=font_size)

    for row in table.rows:
        row.height = Inches(0.2)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    metric_col_width = Inches(1.55)
    data_col_width = Inches(max(1.0, min(1.75, 8.8 / max(len(rows), 1))))

    for row in table.rows:
        row.cells[0].width = metric_col_width

        for cell in row.cells[1:]:
            cell.width = data_col_width


def _set_cell_text(cell, text: str, *, bold: bool = False, size: float = 6.5) -> None:
    cell.text = ""

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


# ======================================================================================
# MASTER DOCX REPORT
# ======================================================================================

def _iter_report_states(final_states):
    if isinstance(final_states, dict):
        final_states = [final_states]

    for item in final_states:
        if not isinstance(item, dict):
            continue

        if item.get("multi") and isinstance(item.get("states"), dict):
            for sub_state in item["states"].values():
                if isinstance(sub_state, dict):
                    yield sub_state

        elif "base_cfg" in item:
            yield item

        elif isinstance(item.get("states"), dict):
            for sub_state in item["states"].values():
                if isinstance(sub_state, dict) and "base_cfg" in sub_state:
                    yield sub_state


def _add_state_to_master_doc(
    *,
    doc,
    state: dict,
    csv_path: Path,
    plot_path: Path | None,
) -> None:
    rows = _read_csv_rows(csv_path)

    if not rows:
        p = doc.add_paragraph()
        p.add_run(f"No table data found for {state.get('name', state.get('pair_tag', 'unknown'))}")
        return

    title_text = state.get("name", state.get("pair_tag", "NEB report"))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(str(title_text))
    title_run.bold = True
    title_run.font.size = Pt(12)

    if plot_path is not None and Path(plot_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        run = p.add_run()
        run.add_picture(str(plot_path), width=Inches(10.35))

    _add_docx_table(
        doc=doc,
        rows=rows,
        compact=True,
    )